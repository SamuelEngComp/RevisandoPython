# TRABALHANDO COM DOCUMENTOS PDF E WORD
# Documentos PDF
import PyPDF2

# Extraindo texto de PDFs
objetoPDF = open('101.pdf','rb')
leitorPDF = PyPDF2.PdfFileReader(objetoPDF)  #lendo o objeto pdf
numeroDePaginas = leitorPDF.numPages #numero de paginas
print(numeroDePaginas)

objetoPagina = leitorPDF.getPage(0)
textoExtraido = objetoPagina.extractText()
print(textoExtraido) ## não funcioou.. acho q é devido a versão do python e da biblioteca

# Descriptografando PDFs
# Criando PDFs
# Copiando páginas
# Rotação de páginas
# Sobrepondo páginas
# Criptografando PDFs
"""NÃO CONSEGUIR UTILIZAR NENHUMA CLASS DO MODULO PyPDF2 -- ACREDITO QUE DEVA SER DEVIDO A VERSÃO DO MEU PYTHON 3.7
QUE NÃO ESTA SENDO COMPTATÍVEL COM A VERSÃO DO MODULO PyPDF2"""

# Documentos Word
import docx

# Lendo documentos Word
doc = docx.Document('Texto-atividade-01.docx')
print(len(doc.paragraphs))

textoDoWord = doc.paragraphs[0].text
textoDoWord01 = doc.paragraphs[1].text
textoDoWord02 = doc.paragraphs[2].text

print(textoDoWord)
print(textoDoWord01)
print(textoDoWord02)

textoRUN = len(doc.paragraphs[1].runs)
print(textoRUN)

depuracaoTexto = doc.paragraphs[0].runs[0].text
print(depuracaoTexto) # Universidade

depuracaoTexto = doc.paragraphs[1].runs[0].text
print(depuracaoTexto) # Universidade

# Obtendo o texto completo de um arquivo .docx
def pegandoTodoTexto(arquivo):
    documento = docx.Document(arquivo)
    textCompleto = []
    for para in documento.paragraphs:
        textCompleto.append(para.text)

    return '\n'.join(textCompleto)

print(pegandoTodoTexto('Texto-atividade-01.docx'))

"""ESSE MODULO É IMPORTANTE, MAS NÃO VEJO MUITA APLICAÇÃO PRÁTICA PRA ELE, ENTÃO VOU PASSAR ELE MAIS RÁPIDO"""
